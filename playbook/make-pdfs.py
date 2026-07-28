#!/usr/bin/env python3
"""
make-pdfs.py - Turn the playbook markdown sheets into printable PDFs and
editable Word documents.

Edit any playbook/*.md, run this, and get:
  playbook/pdf/       print these and tape them up
  playbook/editable/  .docx, opens in Pages or Word, type straight into it

The markdown stays the source of truth. Edits made in Pages live only in that
.docx, so anything meant to stick should go back into the .md (easiest: ask
Claude to make the change, then re-run this).

Usage:
    python3 playbook/make-pdfs.py
    python3 playbook/make-pdfs.py --selfcheck   # verify conversion, no render

Needs: pip install -r requirements.txt && playwright install chromium
"""

from __future__ import annotations

import hashlib
import re
import shutil
import sys
from pathlib import Path

PLAYBOOK_DIR = Path(__file__).resolve().parent
OUT_DIR = PLAYBOOK_DIR / "pdf"
EDITABLE_DIR = PLAYBOOK_DIR / "editable"
BACKUP_DIR = EDITABLE_DIR / "_previous"

# Brand colors. Light background on purpose -- these get printed, and a navy
# page would eat a cartridge per sheet.
NAVY = "#0B1C2D"
GOLD = "#C8922A"
CREAM = "#F5EFD8"

CSS = f"""
@page {{ size: Letter; margin: 0.6in 0.65in; }}
* {{ box-sizing: border-box; }}
body {{
  font-family: "Helvetica Neue", Helvetica, Arial, sans-serif;
  font-size: 10.5pt; line-height: 1.5; color: #1a1a1a; margin: 0;
}}
h1 {{
  font-size: 22pt; color: {NAVY}; margin: 0 0 4pt; letter-spacing: -0.4pt;
  border-bottom: 3px solid {GOLD}; padding-bottom: 7pt;
}}
h2 {{
  font-size: 13.5pt; color: {NAVY}; margin: 20pt 0 6pt;
  border-left: 4px solid {GOLD}; padding-left: 8pt;
  page-break-after: avoid;
}}
h3 {{ font-size: 11.5pt; color: {NAVY}; margin: 13pt 0 4pt; page-break-after: avoid; }}
p {{ margin: 0 0 7pt; }}
ul, ol {{ margin: 0 0 8pt; padding-left: 18pt; }}
li {{ margin-bottom: 3pt; }}
strong {{ color: {NAVY}; }}
hr {{ border: 0; border-top: 1px solid #d8d8d8; margin: 16pt 0; }}
code {{
  font-family: "SF Mono", Menlo, Consolas, monospace; font-size: 9.5pt;
  background: {CREAM}; padding: 1pt 4pt; border-radius: 3px; color: {NAVY};
}}
pre {{
  background: {NAVY}; color: {CREAM}; padding: 9pt 12pt; border-radius: 5px;
  overflow-x: auto; page-break-inside: avoid; margin: 0 0 9pt;
}}
pre code {{ background: none; color: {CREAM}; padding: 0; font-size: 10pt; }}
blockquote {{
  margin: 0 0 9pt; padding: 7pt 12pt; background: #f6f6f4;
  border-left: 3px solid {GOLD}; color: #333; font-style: italic;
}}
blockquote p {{ margin: 0; }}
blockquote p + p {{ margin-top: 7pt; }}
table {{
  width: 100%; border-collapse: collapse; margin: 0 0 11pt; font-size: 9.5pt;
  page-break-inside: avoid;
}}
th {{
  background: {NAVY}; color: {CREAM}; text-align: left; padding: 6pt 8pt;
  font-weight: 600;
}}
td {{ padding: 5pt 8pt; border-bottom: 1px solid #e2e2e2; vertical-align: top; }}
tr:nth-child(even) td {{ background: #fafaf8; }}
/* Blank fill-in cells get a visible ruled line to write on */
td:empty {{ border-bottom: 1px solid {GOLD}; min-width: 90pt; }}
.cb {{
  display: inline-block; width: 10pt; height: 10pt; border: 1.5px solid {NAVY};
  border-radius: 2px; margin-right: 6pt; vertical-align: -1pt;
}}
li.task {{ list-style: none; margin-left: -14pt; margin-bottom: 5pt; }}
.footer {{
  margin-top: 22pt; padding-top: 7pt; border-top: 1px solid #d8d8d8;
  font-size: 8pt; color: #888;
}}
"""

# `- [ ]` task items, after markdown has wrapped them in <li>
TASK_RE = re.compile(r"<li>\s*\[([ xX])\]\s*", re.I)


def md_to_html(md_text: str, title: str) -> str:
    """Convert one markdown sheet into a standalone printable HTML page."""
    import markdown as md_lib

    body = md_lib.markdown(md_text, extensions=["tables", "fenced_code", "sane_lists"])
    body = TASK_RE.sub('<li class="task"><span class="cb"></span>', body)

    return (
        f"<!doctype html><html><head><meta charset='utf-8'>"
        f"<title>{title}</title><style>{CSS}</style></head><body>{body}"
        f"<div class='footer'>Backyard Brew &middot; {title} &middot; "
        f"edit playbook/{title}.md and re-run make-pdfs.py</div>"
        f"</body></html>"
    )


LIST_START = re.compile(r"^\s*(?:[-*+]\s|\d+\.\s)")


def lint(md_text: str) -> list[str]:
    """Warn about a list glued to the line above it.

    Markdown silently swallows the list into the previous paragraph, which is
    the one mistake a non-markdown person makes constantly when editing these
    sheets. Cheap to detect, invisible until it prints wrong.
    """
    warnings = []
    lines = md_text.splitlines()
    for i, line in enumerate(lines[1:], start=2):
        prev = lines[i - 2]
        if LIST_START.match(line) and prev.strip() and not LIST_START.match(prev):
            # A wrapped continuation line inside a list item is fine.
            if prev.startswith(("  ", "\t")):
                continue
            warnings.append(f"  line {i}: list needs a blank line above it")
    return warnings


INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`)")


def _add_runs(paragraph, text: str) -> None:
    """Write text into a docx paragraph, honouring **bold** and `code`."""
    for chunk in INLINE_RE.split(text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            paragraph.add_run(chunk[2:-2]).bold = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            run = paragraph.add_run(chunk[1:-1])
            run.font.name = "Menlo"
        else:
            paragraph.add_run(chunk)


def _docx_source_hash(path: Path) -> str | None:
    """The markdown hash stamped into a .docx when we generated it."""
    if not path.exists():
        return None
    try:
        from docx import Document

        return Document(str(path)).core_properties.comments or None
    except Exception:
        return None


def md_to_docx(md_text: str, out_path: Path, source_hash: str = '') -> None:
    """Write an editable Word document he can open in Pages and print.

    Deliberately a line walker over the constructs these sheets actually use
    (headings, paragraphs, bullets, numbers, checkboxes, tables, quotes,
    fenced code) rather than a general markdown-to-docx converter.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    doc.styles["Normal"].font.name = "Helvetica Neue"
    doc.styles["Normal"].font.size = Pt(11)
    navy = RGBColor(0x0B, 0x1C, 0x2D)

    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped or stripped == "---":
            i += 1
            continue

        # Fenced code
        if stripped.startswith("```"):
            block = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            para = doc.add_paragraph()
            run = para.add_run("\n".join(block))
            run.font.name = "Menlo"
            run.font.size = Pt(10)
            i += 1
            continue

        # Table: a header row followed by a |---| separator
        if stripped.startswith("|") and i + 1 < len(lines) and set(
            lines[i + 1].strip().replace("|", "").replace(" ", "")
        ) <= {"-", ":"} and "-" in lines[i + 1]:
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not (set("".join(cells).replace(" ", "")) <= {"-", ":"} and cells):
                    rows.append(cells)
                i += 1
            width = max(len(r) for r in rows)
            table = doc.add_table(rows=0, cols=width)
            table.style = "Light Grid Accent 1"
            for r_i, cells in enumerate(rows):
                cells += [""] * (width - len(cells))
                row = table.add_row()
                for c_i, text in enumerate(cells):
                    cell_para = row.cells[c_i].paragraphs[0]
                    _add_runs(cell_para, text)
                    if r_i == 0:
                        for run in cell_para.runs:
                            run.bold = True
            doc.add_paragraph()
            continue

        # Headings
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            head = doc.add_heading(level=min(level, 4))
            _add_runs(head, stripped.lstrip("#").strip())
            for run in head.runs:
                run.font.color.rgb = navy
            i += 1
            continue

        # Quote
        if stripped.startswith(">"):
            para = doc.add_paragraph(style="Intense Quote")
            _add_runs(para, stripped.lstrip("> ").strip())
            i += 1
            continue

        # Checkbox, bullet, numbered
        checkbox = re.match(r"[-*+]\s+\[([ xX])\]\s+(.*)", stripped)
        if checkbox:
            para = doc.add_paragraph(style="List Bullet")
            _add_runs(para, "☐  " + checkbox.group(2))
            i += 1
            continue
        if re.match(r"[-*+]\s+", stripped):
            para = doc.add_paragraph(style="List Bullet")
            _add_runs(para, re.sub(r"^[-*+]\s+", "", stripped))
            i += 1
            continue
        if re.match(r"\d+\.\s+", stripped):
            para = doc.add_paragraph(style="List Number")
            _add_runs(para, re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue

        # Paragraph, joining wrapped lines
        buf = [stripped]
        i += 1
        while i < len(lines) and lines[i].strip() and not re.match(
            r"^(#|>|\||```|[-*+]\s|\d+\.\s|---)", lines[i].strip()
        ):
            buf.append(lines[i].strip())
            i += 1
        _add_runs(doc.add_paragraph(), " ".join(buf))

    doc.core_properties.comments = source_hash
    doc.save(str(out_path))


def _selfcheck() -> None:
    """Smallest thing that fails if the conversion breaks."""
    html = md_to_html(
        "# T\n\n- [ ] todo\n\n| a | b |\n|---|---|\n| 1 | |\n\n`code`\n", "T"
    )
    assert '<li class="task"><span class="cb"></span>todo' in html, "checkbox"
    assert "<table>" in html and "<th>a</th>" in html, "table"
    assert "<code>code</code>" in html, "inline code"
    assert "<h1>T</h1>" in html, "heading"
    assert lint("Do this:\n- a\n"), "lint should catch a glued list"
    assert not lint("Do this:\n\n- a\n- b\n"), "lint should pass a spaced list"
    assert not lint("- a\n  wrapped\n- b\n"), "lint should allow continuations"

    import tempfile

    sample = (
        "# Title\n\n## Section\n\nSome **bold** and `code` text that\n"
        "wraps across lines.\n\n- a bullet\n- [ ] a checkbox\n\n"
        "1. numbered\n\n> a quote\n\n| a | b |\n|---|---|\n| 1 | 2 |\n\n"
        "```\nfenced\n```\n"
    )
    with tempfile.TemporaryDirectory() as tmp:
        out = Path(tmp) / "t.docx"
        md_to_docx(sample, out)
        assert out.stat().st_size > 5000, "docx looks empty"
        from docx import Document

        doc = Document(str(out))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Title" in text and "Section" in text, "docx headings"
        assert "bold" in text and "wraps across lines" in text, "docx paragraph"
        assert "☐" in text, "docx checkbox"
        assert len(doc.tables) == 1, "docx table"
        assert doc.tables[0].rows[0].cells[0].text == "a", "docx table header"
    print("selfcheck ok")


def main() -> int:
    if "--selfcheck" in sys.argv:
        _selfcheck()
        return 0

    try:
        import markdown  # noqa: F401
        from playwright.sync_api import sync_playwright
    except ImportError as exc:
        print(f"Missing dependency: {exc.name}")
        print("Run:  pip install -r requirements.txt && playwright install chromium")
        return 1

    sheets = sorted(PLAYBOOK_DIR.glob("[0-9]-*.md"))
    if not sheets:
        print(f"No playbook sheets found in {PLAYBOOK_DIR}")
        return 1

    OUT_DIR.mkdir(exist_ok=True)
    EDITABLE_DIR.mkdir(exist_ok=True)
    replaced: list[str] = []

    with sync_playwright() as p:
        browser = p.chromium.launch()
        page = browser.new_page()
        for sheet in sheets:
            text = sheet.read_text(encoding="utf-8")
            for warning in lint(text):
                print(f"  ! {sheet.name}{warning}")
            html = md_to_html(text, sheet.stem)
            page.set_content(html, wait_until="load")
            page.pdf(
                path=str(OUT_DIR / f"{sheet.stem}.pdf"),
                format="Letter",
                print_background=True,
            )
            # Only rewrite a .docx when its source markdown actually changed.
            # Someone may have typed into it in Pages (the Facebook groups
            # table is designed to be filled in by hand), and silently
            # overwriting that is real lost work.
            docx_path = EDITABLE_DIR / f"{sheet.stem}.docx"
            source_hash = hashlib.sha256(text.encode("utf-8")).hexdigest()[:16]

            if "--force" not in sys.argv and _docx_source_hash(docx_path) == source_hash:
                print(f"  {sheet.name}  ->  pdf/  (.docx already current)")
            else:
                if docx_path.exists():
                    # The markdown moved on. Keep the old copy in case it held
                    # hand edits, rather than assuming it did not.
                    BACKUP_DIR.mkdir(exist_ok=True)
                    shutil.copy2(docx_path, BACKUP_DIR / docx_path.name)
                    replaced.append(sheet.stem)
                md_to_docx(text, docx_path, source_hash)
                print(f"  {sheet.name}  ->  pdf/ + editable/")
        browser.close()

    if replaced:
        print(
            f"\nRewrote {len(replaced)} .docx file(s) because the sheet changed:"
            f"\n  {', '.join(replaced)}"
            f"\nThe previous versions are in {BACKUP_DIR.name}/ if anything was"
            "\ntyped into them by hand."
        )

    print(f"\nDone. {len(sheets)} sheets. Print from pdf/, edit in editable/")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
